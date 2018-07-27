# -*- coding: utf-8 -*-
"""
Created on Mon Jul 23 20:25:09 2018

@author: Santosh
"""
"""
from load_words import bad_words_list_final as words
def profanityFilter(text):
    brokenStr1 = text.split()
    badWordMask = '!@#$%!@#$%^~!@%^~@#$%!@#$%^~!'
    for word in brokenStr1:
        if word in words:
            print (word + ' <--Bad word!')
            text = text.replace(word,badWordMask[:len(word)])
            #print new

    return text


def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text=profanityFilter(inline[i].text)
                    #text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                
                
#print (profanityFilter("this thing motherfuckin' sucks sucks fucking stuff"))
from docx import Document
filename = "Hello12.docx"
doc = Document(filename)
for word in words:
    
"""
with open('text.txt', 'r') as myfile:
    data=myfile.read()
from profanityfilter import ProfanityFilter
pf = ProfanityFilter()
#print(pf.censor("That's bullshit!"))
from load_words import bad_words_list_final as words
pf.append_words(words)
filename = "Hello12.docx"
#doc = Document(filename)
#for p in doc.paragraphs:
    #inline=p.runs
    #for i in range(len(inline)):
        #y=inline[i].text
        #print(pf.censor(y))
        
        
data_filter=pf.censor(data)
f = open('filter.txt','w')
f.write(data_filter)
f.close()
#print(pf.censor('Unbuttoned, butt, buttoned'))