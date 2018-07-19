# -*- coding: utf-8 -*-
"""
Created on Thu Jul 19 14:18:25 2018

@author: Santosh
"""


import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, " "+regex+" " , replace)

from load_words import bad_words_list_final

filename = "Hello12.docx"
doc = Document(filename)
replacement=[]
for i in range(len(bad_words_list_final)):
    replacement.append(bad_words_list_final[i][0]+(len(bad_words_list_final[i])-1)*'*')
Dictionary = dict(zip(bad_words_list_final, replacement))
#Dictionary={bad_words_list_final,replacement}

for word, replacement in Dictionary.items():
    word_re=re.compile(" "+word+" ")
    docx_replace_regex(doc, word_re , " "+replacement+" ")

#regex1 = re.compile(r"Thnks")
#replace1 = r"Thanks"

#docx_replace_regex(doc, regex1 , replace1)
doc.save('Hello12_filtered.docx')